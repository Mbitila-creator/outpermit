from io import BytesIO

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

from .qr_cards import render_participant_qr_only


NAVY = RGBColor(23, 54, 93)
MUTED = RGBColor(100, 116, 139)
BORDER = "B9C7D6"
TABLE_WIDTH_DXA = 10658
QR_COLUMN_DXA = 3118
TEXT_COLUMN_DXA = TABLE_WIDTH_DXA - QR_COLUMN_DXA


def _set_font(run, *, size, bold=False, color=None):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _set_cell_margins(cell, *, top, start, bottom, end):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        element = tc_mar.find(qn(f"w:{margin_name}"))
        if element is None:
            element = OxmlElement(f"w:{margin_name}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_card_table_geometry(table):
    table.autofit = False
    table.alignment = 1
    table.columns[0].width = Mm(55)
    table.columns[1].width = Mm(133)
    for cell, width in zip(table.rows[0].cells, (QR_COLUMN_DXA, TEXT_COLUMN_DXA)):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = tc_pr.first_child_found_in("w:tcW")
        tc_w.set(qn("w:w"), str(width))
        tc_w.set(qn("w:type"), "dxa")

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in (QR_COLUMN_DXA, TEXT_COLUMN_DXA):
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "start", "bottom", "end"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), BORDER)
        borders.append(border)
    for edge in ("insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "nil")
        borders.append(border)
    tbl_pr.append(borders)


def _add_text_paragraph(cell, text, *, size, bold=False, color=None, after=0):
    paragraph = cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text or "—")
    _set_font(run, size=size, bold=bold, color=color)
    return paragraph


def _add_participant_card(document, participant, verification_url):
    table = document.add_table(rows=1, cols=2)
    _set_card_table_geometry(table)
    row = table.rows[0]
    row.height = Mm(86)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    cant_split = OxmlElement("w:cantSplit")
    row._tr.get_or_add_trPr().append(cant_split)

    qr_cell, text_cell = row.cells
    qr_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    text_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(qr_cell, top=220, start=360, bottom=220, end=220)
    _set_cell_margins(text_cell, top=220, start=180, bottom=220, end=360)

    qr_paragraph = qr_cell.paragraphs[0]
    qr_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    qr_paragraph.paragraph_format.space_before = Pt(0)
    qr_paragraph.paragraph_format.space_after = Pt(0)
    qr_run = qr_paragraph.add_run()
    qr_run.add_picture(
        BytesIO(render_participant_qr_only(verification_url)),
        width=Mm(43),
        height=Mm(43),
    )

    publications = getattr(participant, "active_publications", None)
    publication_count = (
        len(publications)
        if publications is not None
        else participant.publications.filter(is_active=True).count()
    )
    publication_label = "publication" if publication_count == 1 else "publications"
    identity = f"{participant.event.code} · {publication_count} {publication_label}"
    _add_text_paragraph(
        text_cell,
        identity,
        size=9,
        bold=True,
        color=NAVY,
        after=7,
    )
    _add_text_paragraph(
        text_cell,
        participant.full_name,
        size=15.5,
        bold=True,
        after=5,
    )
    _add_text_paragraph(
        text_cell,
        participant.institution,
        size=9.5,
        after=7,
    )
    _add_text_paragraph(
        text_cell,
        "Scan to view the verified researcher record",
        size=9,
        color=MUTED,
    )


def build_editable_participant_cards(participants, verification_url_for):
    """Build one editable Word file with two cards per A4 page."""
    participants = list(participants)
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(10)
    section.right_margin = Mm(10)
    section.bottom_margin = Mm(10)
    section.left_margin = Mm(10)
    section.header_distance = Mm(5)
    section.footer_distance = Mm(5)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    document.core_properties.title = "Special Event QR cards"
    document.core_properties.subject = "Editable participant QR cards"

    for index, participant in enumerate(participants):
        _add_participant_card(
            document,
            participant,
            verification_url_for(participant),
        )
        is_last = index == len(participants) - 1
        if is_last:
            continue
        if (index + 1) % 2 == 0:
            document.add_page_break()
        else:
            spacer = document.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(9)
            spacer.add_run("")

    output = BytesIO()
    document.save(output)
    return output.getvalue()

